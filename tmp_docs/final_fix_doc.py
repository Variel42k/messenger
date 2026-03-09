from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = Path(r'e:\git\messenger\tmp_docs\target.docx')
doc = Document(str(DOC_PATH))

STYLE_NAMES = [s.name for s in doc.styles]
H0 = 'САМзагаловок типа ВВЕДЕНИЕ' if 'САМзагаловок типа ВВЕДЕНИЕ' in STYLE_NAMES else ('САМзаголовок 1-ого уровня' if 'САМзаголовок 1-ого уровня' in STYLE_NAMES else 'Heading 1')
BODY = 'САМосновной текст' if 'САМосновной текст' in STYLE_NAMES else 'Normal'


def p_text(p):
    return ' '.join((p.text or '').split())


def set_run_font(par, bold=False):
    for r in par.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.bold = bold


def normalize_paragraph(par, heading=False, front=False):
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if heading or front:
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.0
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER if front else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(par, bold=True)
    else:
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = 1.5
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_run_font(par, bold=False)


def find_idx_starts(prefix):
    for i, p in enumerate(doc.paragraphs):
        if p_text(p).startswith(prefix):
            return i
    return None


def find_idx_exact(text):
    for i, p in enumerate(doc.paragraphs):
        if p_text(p) == text:
            return i
    return None


def insert_before_idx(idx, text, style_name=None):
    if idx is None:
        p = doc.add_paragraph(text)
    else:
        p = doc.paragraphs[idx].insert_paragraph_before(text)
    if style_name and style_name in STYLE_NAMES:
        p.style = doc.styles[style_name]
    return p


def remove_paragraph(par):
    el = par._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

# 1) Rename headings to exact required wording
for p in doc.paragraphs:
    t = p_text(p)
    if t.startswith('2.4.2 Цели разработки веб-приложения мессенджера'):
        p.text = '2.4.2 Цели разработки системы'
        normalize_paragraph(p, heading=True)
    elif t.startswith('2.4.8 Описание экранных форм'):
        p.text = '2.4.8 Проверка прототипа по сценариям'
        normalize_paragraph(p, heading=True)

# 2) Add missing section "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ"
abbr_heading = 'ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ'
if find_idx_exact(abbr_heading) is None:
    idx_intro = find_idx_exact('ВВЕДЕНИЕ')
    if idx_intro is None:
        idx_intro = find_idx_starts('1 Постановка задачи')
    p_h = insert_before_idx(idx_intro, abbr_heading, H0)
    normalize_paragraph(p_h, heading=True, front=True)

    abbr_lines = [
        'ИС - информационная система;',
        'БД - база данных;',
        'API - программный интерфейс приложения;',
        'UI - пользовательский интерфейс;',
        'UX - пользовательский опыт;',
        'JWT - JSON Web Token.',
    ]
    # insert in reverse so final order is correct
    for line in reversed(abbr_lines):
        p_l = insert_before_idx(idx_intro, line, BODY)
        normalize_paragraph(p_l, heading=False)

# 3) Remove corrupted artifact paragraphs at end (question-mark garbage)
#    Keep valid appendix/code content untouched.
for i in range(len(doc.paragraphs) - 1, max(-1, len(doc.paragraphs) - 30), -1):
    p = doc.paragraphs[i]
    t = p_text(p)
    if not t:
        continue
    t_no_q = t.replace('?', '').replace(' ', '').replace('.', '').replace(',', '').replace(':', '').replace(';', '').replace('(', '').replace(')', '').replace('-', '')
    if t_no_q == '' and '?' in t:
        remove_paragraph(p)

# also remove likely broken tail paragraphs if they are mostly '?'
for i in range(len(doc.paragraphs) - 1, max(-1, len(doc.paragraphs) - 20), -1):
    p = doc.paragraphs[i]
    t = p_text(p)
    if not t:
        continue
    q_ratio = t.count('?') / max(1, len(t))
    if q_ratio > 0.5:
        remove_paragraph(p)

# 4) Ensure appendix heading text exactly matches requirement
idx_app = find_idx_starts('ПРИЛОЖЕНИЕ А')
if idx_app is not None:
    doc.paragraphs[idx_app].text = 'ПРИЛОЖЕНИЕ А (обязательное) Исходный текст программы'
    if H0 in STYLE_NAMES:
        doc.paragraphs[idx_app].style = doc.styles[H0]
    normalize_paragraph(doc.paragraphs[idx_app], heading=True, front=True)

doc.save(str(DOC_PATH))
print('final_fix_done')
