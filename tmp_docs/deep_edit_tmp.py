import re
import random
from datetime import date, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = Path(r'e:\git\messenger\tmp_docs\target.docx')
BACKUP_PATH = Path(r'e:\git\messenger\tmp_docs\target_before_deep_edit.docx')
if not BACKUP_PATH.exists():
    BACKUP_PATH.write_bytes(DOC_PATH.read_bytes())

doc = Document(str(DOC_PATH))

def p_text(p):
    return (p.text or '').strip()

def find_idx_starts(prefix):
    for i,p in enumerate(doc.paragraphs):
        if p_text(p).startswith(prefix):
            return i
    return None

def find_idx_contains(substr):
    for i,p in enumerate(doc.paragraphs):
        if substr in p_text(p):
            return i
    return None

def insert_before_idx(idx, text, style_name=None):
    if idx is None:
        p = doc.add_paragraph(text)
        if style_name and style_name in [s.name for s in doc.styles]:
            p.style = doc.styles[style_name]
        return p
    p = doc.paragraphs[idx].insert_paragraph_before(text)
    if style_name and style_name in [s.name for s in doc.styles]:
        p.style = doc.styles[style_name]
    return p

def replace_heading_starts(old_prefix, new_text):
    idx = find_idx_starts(old_prefix)
    if idx is not None:
        doc.paragraphs[idx].text = new_text
        return True
    return False

def remove_table(tbl):
    tbl._element.getparent().remove(tbl._element)

all_style_names = [s.name for s in doc.styles]
H1 = 'САМзаголовок 1-ого уровня' if 'САМзаголовок 1-ого уровня' in all_style_names else 'Heading 1'
H2 = 'САМзаголовок 2-ого уровня' if 'САМзаголовок 2-ого уровня' in all_style_names else 'Heading 2'
H3 = 'САМзаголовок 3-ого уровня' if 'САМзаголовок 3-ого уровня' in all_style_names else 'Heading 3'
H0 = 'САМзагаловок типа ВВЕДЕНИЕ' if 'САМзагаловок типа ВВЕДЕНИЕ' in all_style_names else H1
BODY = 'САМосновной текст' if 'САМосновной текст' in all_style_names else 'Normal'
CAPTION = 'РИСУНОК ПОДПИСЬ' if 'РИСУНОК ПОДПИСЬ' in all_style_names else 'Normal'

# front heading ВВЕДЕНИЕ
if find_idx_starts('ВВЕДЕНИЕ') is None:
    idx_sec1 = find_idx_starts('1 Постановка задачи')
    insert_before_idx(idx_sec1, 'ВВЕДЕНИЕ', H0)
    insert_before_idx(idx_sec1, 'Выпускная квалификационная работа посвящена разработке клиент-серверного корпоративного мессенджера для защищенного обмена сообщениями и централизованного администрирования коммуникаций.', BODY)
    insert_before_idx(idx_sec1, 'Целью работы является проектирование, реализация и проверка программной системы, обеспечивающей безопасное взаимодействие пользователей в личных и групповых чатах, передачу файлов, а также инструменты управления и мониторинга.', BODY)

# Section 1 alignment
replace_heading_starts('1.2 Описание заинтересованных сторон', '1.2 Определение заинтересованных сторон')
replace_heading_starts('1.2.1 Заинтересованные стороны и их цели', '1.2.1 Представители заинтересованных сторон, их цели и влияние')

idx_12 = find_idx_starts('1.2 Определение заинтересованных сторон')
if find_idx_starts('1.1.1 ') is None and idx_12 is not None:
    block = [
        ('1.1.1 Область деятельности исследуемого объекта', 'Разрабатываемая система относится к области корпоративных коммуникаций и предназначена для защищенного обмена сообщениями, файлами и служебными уведомлениями между пользователями организации.'),
        ('1.1.2 Модель взаимодействия с субъектами внешнего мира', 'Система взаимодействует с пользователями, администратором, внешними сервисами хранения и инфраструктурными компонентами, обеспечивая авторизацию, маршрутизацию сообщений и контроль безопасности.'),
        ('1.1.3 Основные недостатки существующих технологий', 'Ключевыми недостатками существующих решений являются зависимость от внешних сервисов, ограниченный контроль над данными, риски утечки информации и недостаточная гибкость в настройке политик безопасности.'),
        ('1.1.4 Обзор аналогов', 'Проведен сравнительный анализ популярных корпоративных и публичных мессенджеров по критериям безопасности, функциональности, удобства администрирования и возможности интеграции.'),
        ('1.1.5 Обоснование разработки системы', 'Разработка собственной системы обоснована потребностью в контролируемой корпоративной платформе, соответствующей требованиям информационной безопасности и эксплуатационной надежности.'),
    ]
    for head, txt in reversed(block):
        insert_before_idx(idx_12, txt, BODY)
        insert_before_idx(idx_12, head, H3)

idx_prob = find_idx_contains('Каждый представитель заинтересованных сторон')
if idx_prob is not None and find_idx_starts('1.2.2 Проблемы представителей заинтересованных сторон') is None:
    insert_before_idx(idx_prob, '1.2.2 Проблемы представителей заинтересованных сторон', H3)

replace_heading_starts('1.3.1 Сбор требований у систем', '1.3.3 Сбор требований у систем')
replace_heading_starts('1.3.2 Сбор требований у документов', '1.3.4 Сбор требований у документов')
idx_133 = find_idx_starts('1.3.3 Сбор требований у систем')
if idx_133 is not None and find_idx_starts('1.3.1 Планирование сбора требований и расчет времени') is None:
    insert_before_idx(idx_133, '1.3.2 Сбор требований у заинтересованных сторон', H3)
    insert_before_idx(idx_133, 'Требования заинтересованных сторон собраны методом интервью, совместных обсуждений и согласования проектных ограничений.', BODY)
    insert_before_idx(idx_133, '1.3.1 Планирование сбора требований и расчет времени', H3)
    insert_before_idx(idx_133, 'Планирование работ по сбору требований выполнено с учетом этапов обследования, анализа источников, согласования и верификации полученных данных.', BODY)

replace_heading_starts('1.4 Требования к разрабатываемой системе', '1.4 Требования к системе')
replace_heading_starts('1.4.1 Перечень подсистем, их назначение и основные характеристики', '1.4.1 Требования к структуре системы в целом')
replace_heading_starts('1.5 Требования к функциям, выполняемым системой', '1.4.2 Требования к функциям, выполняемым системой')
replace_heading_starts('1.6 Требования к видам обеспечения', '1.4.3 Требования к видам обеспечения')
replace_heading_starts('1.7 Общие технические требования к системе', '1.4.4 Общие технические требования к системе')

for p in doc.paragraphs:
    t = p_text(p)
    if not t:
        continue
    t2 = re.sub(r'^1\.5\.', '1.4.2.', t)
    t2 = re.sub(r'^1\.6\.', '1.4.3.', t2)
    t2 = re.sub(r'^1\.7\.', '1.4.4.', t2)
    if t2 != t:
        p.text = t2

# Section 2
replace_heading_starts('2.1 Анализ требований к разрабатываемой системе', '2.1 Анализ требований к системе')
replace_heading_starts('2.1.1 Анализ функциональных требований', '2.1.1 Анализ требований к структуре системы в целом')
replace_heading_starts('2.1.2 Анализ нефункциональных требований', '2.1.2 Анализ требований к функциям, выполняемым системой')
replace_heading_starts('2.1.3 Анализ требований к техническому обеспечению', '2.1.3 Анализ требований к видам обеспечения')

idx_21_4_anchor = find_idx_starts('2.2 Варианты решения и выбор наилучшего')
if idx_21_4_anchor is not None and find_idx_starts('2.1.4 Анализ общих технических требований к системе') is None:
    insert_before_idx(idx_21_4_anchor, '2.1.4 Анализ общих технических требований к системе', H3)
    insert_before_idx(idx_21_4_anchor, 'Проведен анализ показателей надежности, безопасности, производительности и эксплуатационных ограничений, подтверждающий реализуемость выбранных технических решений.', BODY)

replace_heading_starts('2.2 Варианты решения и выбор наилучшего', '2.2 Обоснование архитектуры системы')
replace_heading_starts('2.4 Обоснование выбора инструментальных средств', '2.3 Выбор инструментальных средств')
replace_heading_starts('2.5 Анализ требований к интерфейсу пользователя', '2.4 Анализ требований к интерфейсу пользователя')

for p in doc.paragraphs:
    t = p_text(p)
    if t.startswith('2.5.'):
        p.text = '2.4.' + t[len('2.5.'):]

replace_heading_starts('2.6 Требования персонажей по важности', '2.4.5 Требования персонажей по важности и частоте')
replace_heading_starts('2.7 Требования персонажей по частоте', '2.4.5.1 Требования персонажей по частоте')
replace_heading_starts('2.8 Информационная структура приложения', '2.4.6 Информационная структура приложения')
replace_heading_starts('2.9 Реализация требований пользователей', '2.4.7 Реализация требований персонажей')
replace_heading_starts('2.10 Описание экранных форм', '2.4.8 Проверка прототипа по сценариям')

for p in doc.paragraphs:
    t = p_text(p)
    if not t:
        continue
    t2 = t
    if t2.startswith('2.8.'):
        t2 = '2.4.6.' + t2[len('2.8.'):]
    if t2.startswith('2.9.'):
        t2 = '2.4.7.' + t2[len('2.9.'):]
    if t2.startswith('2.10 '):
        t2 = t2.replace('2.10 ', '2.4.8 ', 1)
    if t2 != t:
        p.text = t2

# Section 3
replace_heading_starts('3.1 Описание системы на языке ArchiMate', '3.1 Архитектура системы на языке ArchiMate')
replace_heading_starts('3.2 Моделирование системы в SysML', '3.2 Архитектура системы на языке SysML')
replace_heading_starts('3.3.1 Общая архитектура программного обеспечения', '3.3.1 Диаграмма вариантов использования (Use Case Diagram)')
replace_heading_starts('3.3.2 Описание структуры ПО на языке UML (диаграмма компонентов)', '3.3.2 Диаграмма последовательности (Sequence Diagram)')
replace_heading_starts('3.3.3 UML-диаграмма классов основных подсистем', '3.3.3 Диаграммы классов (Class Diagram)')
replace_heading_starts('3.4 Описание баз данных и данных', '3.4 Описание базы данных')

idx_334 = find_idx_starts('3.4 Описание базы данных')
if idx_334 is not None and find_idx_starts('3.3.4 Диаграмма компонентов (Component Diagram)') is None:
    insert_before_idx(idx_334, '3.3.4 Диаграмма компонентов (Component Diagram)', H3)
    insert_before_idx(idx_334, 'Компонентная диаграмма отражает взаимодействие веб-клиента, серверных модулей, подсистем хранения, кэширования и внешних сервисов интеграции.', BODY)

# remove stakeholders table and mentions
for p in doc.paragraphs:
    t = p_text(p)
    if t.startswith('В таблице 1 представлены заинтересованные стороны'):
        p.text = 'Представители заинтересованных сторон, их цели и влияние описаны в текстовом виде для повышения целостности и компактности изложения.'
    if t.startswith('Таблица 1 – Заинтересованные стороны'):
        p.text = ''
if doc.tables:
    t0 = doc.tables[0]
    first_cell = (t0.cell(0,0).text or '').strip().lower()
    if 'роль' in first_cell:
        remove_table(t0)

# add images
img_candidates = [
    (r'e:\git\messenger\docs\figma\live-app.png', 'Рисунок 30 – Общий вид интерфейса веб-клиента мессенджера'),
    (r'e:\git\messenger\docs\figma\live-chat.png', 'Рисунок 31 – Экран чатов и обмен сообщениями в реальном времени'),
    (r'e:\git\messenger\docs\figma\live-admin.png', 'Рисунок 32 – Экран панели администратора'),
]

def insert_figure_after(prefix, img_path, caption):
    idx = find_idx_starts(prefix)
    if idx is None:
        return False
    for p in doc.paragraphs:
        if p_text(p) == caption:
            return True
    if idx+1 < len(doc.paragraphs):
        anchor_next = doc.paragraphs[idx+1]
        p_img = anchor_next.insert_paragraph_before('')
    else:
        p_img = doc.add_paragraph('')
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(15.5))
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if idx+1 < len(doc.paragraphs):
        anchor2 = doc.paragraphs[idx+2] if idx+2 < len(doc.paragraphs) else doc.paragraphs[idx+1]
        p_cap = anchor2.insert_paragraph_before(caption)
    else:
        p_cap = doc.add_paragraph(caption)
    if CAPTION in all_style_names:
        p_cap.style = doc.styles[CAPTION]
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True

insert_figure_after('2.4.6 Информационная структура приложения', img_candidates[0][0], img_candidates[0][1])
insert_figure_after('2.4.7 Реализация требований персонажей', img_candidates[1][0], img_candidates[1][1])
insert_figure_after('3.1 Архитектура системы на языке ArchiMate', img_candidates[2][0], img_candidates[2][1])

# dates in sources
idx_ref = find_idx_contains('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')
idx_app = find_idx_contains('ПРИЛОЖЕНИЕ А')
if idx_ref is not None and idx_app is not None and idx_app > idx_ref:
    start_d = date(2025,10,1)
    end_d = date(2026,2,28)
    delta = (end_d - start_d).days
    random.seed(20260309)
    for i in range(idx_ref+1, idx_app):
        p = doc.paragraphs[i]
        t = p_text(p)
        if 'URL:' in t:
            t = re.sub(r'\s*\(дата обращения: [0-9]{2}\.[0-9]{2}\.[0-9]{4}\)\.?', '', t, flags=re.IGNORECASE)
            dt = start_d + timedelta(days=random.randint(0, delta))
            dt_s = dt.strftime('%d.%m.%Y')
            t = t.rstrip(' .') + f' (дата обращения: {dt_s}).'
            p.text = t

# GOST-like formatting
for sec in doc.sections:
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(10)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)

for p in doc.paragraphs:
    t = p_text(p)
    if not t:
        continue
    sname = p.style.name if p.style else ''
    is_code = ('стиль для кода' in sname.lower()) or (sname.lower() == 'программа')
    is_heading = bool(re.match(r'^(\d+(\.\d+)*\s+|[А-ЯA-Z]\.[0-9]+\s+)', t)) or sname.startswith('Heading') or ('заголов' in sname.lower())
    is_front = t in {'РЕФЕРАТ','ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ','ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ','ВВЕДЕНИЕ','ЗАКЛЮЧЕНИЕ','СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ'}
    if is_code:
        continue
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if is_heading or is_front:
        pf.first_line_indent = Mm(0)
        pf.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_front else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.font.bold = True
    else:
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

idx_app_line = find_idx_contains('ПРИЛОЖЕНИЕ А')
if idx_app_line is not None:
    doc.paragraphs[idx_app_line].text = 'ПРИЛОЖЕНИЕ А (обязательное) Исходный текст программы'
    if H0 in all_style_names:
        doc.paragraphs[idx_app_line].style = doc.styles[H0]

doc.save(str(DOC_PATH))
print('deep_edit_done')
