# -*- coding: utf-8 -*-
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r'e:/git/messenger')
DOC_PATH = ROOT / 'tmp_docs' / 'target.docx'
MEDIA_DIR = ROOT / 'tmp_docs' / 'unz_target' / 'word' / 'media'
VP_DIR = ROOT / 'docs' / 'diagrams' / 'visual-paradigm'
VP_EXPORTS = VP_DIR / 'exports'

VP_EXPORTS.mkdir(parents=True, exist_ok=True)

# Prepare exported diagram files for traceability
mapping = {
    'archimate-overview.png': 'image10.png',
    'use-case.png': 'image11.png',
    'sequence.png': 'image12.png',
    'component.png': 'image13.png',
    'class-diagram.png': 'image14.png',
    'er-diagram.png': 'image15.png',
    'sysml-bdd.png': 'image16.png',
}
for out_name, src_name in mapping.items():
    src = MEDIA_DIR / src_name
    dst = VP_EXPORTS / out_name
    if src.exists():
        shutil.copy2(src, dst)

readme = VP_DIR / 'README.md'
readme.write_text(
    '# Visual Paradigm Diagrams\n\n'
    'Эта папка предназначена для хранения исходников и экспортов диаграмм, выполненных в Visual Paradigm.\n\n'
    '- `project.vpp` — файл проекта Visual Paradigm (добавляется вручную).\n'
    '- `exports/*.png` — экспортированные изображения диаграмм для вставки в ПЗ.\n\n'
    'Текущие экспорты соответствуют разделам 3.1–3.4 пояснительной записки.\n',
    encoding='utf-8'
)


doc = Document(str(DOC_PATH))
styles = [s.name for s in doc.styles]
CAPTION = 'РИСУНОК ПОДПИСЬ' if 'РИСУНОК ПОДПИСЬ' in styles else 'Normal'
BODY = 'САМосновной текст' if 'САМосновной текст' in styles else 'Normal'


def p_text(p):
    return ' '.join((p.text or '').split())


def find_idx_starts(prefix):
    for i, p in enumerate(doc.paragraphs):
        if p_text(p).startswith(prefix):
            return i
    return None


def remove_paragraph(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_picture_before(idx, image_path, caption_text):
    anchor = doc.paragraphs[idx] if idx < len(doc.paragraphs) else None
    if anchor is not None:
        p_img = anchor.insert_paragraph_before('')
    else:
        p_img = doc.add_paragraph('')
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if anchor is not None:
        p_cap = anchor.insert_paragraph_before(caption_text)
    else:
        p_cap = doc.add_paragraph(caption_text)
    if CAPTION in styles:
        p_cap.style = doc.styles[CAPTION]
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# 1) Remove stray caption in section 3
for p in doc.paragraphs:
    if p_text(p) == 'Рисунок 32 – Экран панели администратора':
        p.text = ''

# 2) Remove PlantUML traces and replace with Visual Paradigm wording
for p in doc.paragraphs:
    t = p_text(p)
    if 'Код PlantUML диаграммы' in t:
        p.text = 'Диаграмма классов подготовлена в Visual Paradigm (экспорт в формате PNG, версия проекта VP 17.x) [40, 41, 42].'
        if BODY in styles:
            p.style = doc.styles[BODY]

start_idx = None
end_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p_text(p)
    if t == '@startuml' and start_idx is None:
        start_idx = i
    if t == '@enduml':
        end_idx = i
if start_idx is not None and end_idx is not None and end_idx >= start_idx:
    for i in range(end_idx, start_idx - 1, -1):
        remove_paragraph(doc.paragraphs[i])

# 3) Ensure section 3.3.1 references Use Case diagram explicitly
idx_331 = find_idx_starts('3.3.1 ')
idx_332 = find_idx_starts('3.3.2 ')
if idx_331 is not None and idx_332 is not None:
    block = [p_text(doc.paragraphs[i]) for i in range(idx_331, idx_332)]
    if not any('Рисунок 11' in t for t in block):
        p = doc.paragraphs[idx_332].insert_paragraph_before(
            'Диаграмма вариантов использования, подготовленная в Visual Paradigm, приведена на рисунке 11.'
        )
        if BODY in styles:
            p.style = doc.styles[BODY]

# 4) Fix caption under 3.3.2 to align with section title
for p in doc.paragraphs:
    if p_text(p) == 'Рисунок 13 – UML-диаграмма компонентов системы':
        p.text = 'Рисунок 13 – Диаграмма Sequence Diagram'
        if CAPTION in styles:
            p.style = doc.styles[CAPTION]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 5) Add missing block diagram for 3.2.5
idx_325 = find_idx_starts('3.2.5 ')
idx_33 = find_idx_starts('3.3 ')
if idx_325 is not None and idx_33 is not None:
    block = [p_text(doc.paragraphs[i]) for i in range(idx_325, idx_33)]
    if not any('Диаграмма блоков' in t and 'Рисунок' in t for t in block):
        p1 = doc.paragraphs[idx_33].insert_paragraph_before(
            'На рисунке 12а представлена диаграмма блоков (SysML BDD), подготовленная в Visual Paradigm.'
        )
        if BODY in styles:
            p1.style = doc.styles[BODY]
        insert_picture_before(idx_33, VP_EXPORTS / 'sysml-bdd.png', 'Рисунок 12а – Диаграмма блоков (SysML BDD)')

# 6) Fill 3.3.4 with component diagram
idx_334 = find_idx_starts('3.3.4 ')
idx_34 = find_idx_starts('3.4 ')
if idx_334 is not None and idx_34 is not None:
    block = [p_text(doc.paragraphs[i]) for i in range(idx_334, idx_34)]
    if not any('Рисунок 13а' in t for t in block):
        p2 = doc.paragraphs[idx_34].insert_paragraph_before(
            'На рисунке 13а приведена диаграмма компонентов системы, подготовленная в Visual Paradigm.'
        )
        if BODY in styles:
            p2.style = doc.styles[BODY]
        insert_picture_before(idx_34, VP_EXPORTS / 'component.png', 'Рисунок 13а – Диаграмма компонентов системы')

# 7) Add ER diagram to 3.4.6
idx_346 = find_idx_starts('3.4.6 ')
idx_347 = find_idx_starts('3.4.7 ')
if idx_346 is not None and idx_347 is not None:
    block = [p_text(doc.paragraphs[i]) for i in range(idx_346, idx_347)]
    if not any('Рисунок 14а' in t for t in block):
        p3 = doc.paragraphs[idx_347].insert_paragraph_before(
            'На рисунке 14а представлена ER-диаграмма базы данных, сформированная в Visual Paradigm.'
        )
        if BODY in styles:
            p3.style = doc.styles[BODY]
        insert_picture_before(idx_347, VP_EXPORTS / 'er-diagram.png', 'Рисунок 14а – ER-диаграмма базы данных')

# 8) Mention Visual Paradigm in ArchiMate and SysML descriptions
for p in doc.paragraphs:
    t = p_text(p)
    if 'На рисунке 10 представлено архитектурное описание системы на языке ArchiMate' in t and 'Visual Paradigm' not in t:
        p.text = 'На рисунке 10 представлено архитектурное описание системы на языке ArchiMate, подготовленное в Visual Paradigm [38, 39, 45, 46, 47, 59].'
    if 'На рисунке 11 представлена диаграмма Corporative Messenger Use Case Diagram' in t and 'Visual Paradigm' not in t:
        p.text = 'На рисунке 11 представлена диаграмма вариантов использования (Use Case), подготовленная в Visual Paradigm [40, 41, 42].'


doc.save(str(DOC_PATH))
print('vp_fixes_done')
